import os
import base64
import time
import io
from io import BytesIO
from flask_cors import cross_origin
from flask import Flask, request, jsonify
from pydub import AudioSegment
import spacy
from spacy.lang.nl.stop_words import STOP_WORDS
from string import punctuation
from heapq import nlargest
from openai import OpenAI
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.prompts import PromptTemplate
from langchain.chat_models import ChatOpenAI
from langchain.chains import LLMChain
from langchain.vectorstores import Qdrant
from langchain.chains import RetrievalQA
from qdrant_client import QdrantClient
import logging
import json
from google.cloud import storage
from google.cloud import tasks_v2
from google.cloud import firestore
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from docx import Document

CLOUD_BUCKET_NAME = "stt-noa"
PROJECT = 'stt2langchain1'

client = tasks_v2.CloudTasksClient()
project = PROJECT
queue = 'NOA'
location = 'europe-west1'
parent = client.queue_path(project, location, queue)
nlp = spacy.load('nl_core_news_lg')

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)
global_log_fields = {}
def log_message(message, severity="NOTICE", **additional_fields):
    entry = {
        "message": message,
        "severity": severity,
        **global_log_fields,
        **additional_fields,
    }
    print(json.dumps(entry))

app = Flask(__name__)

def process_message(message):
#process the message from the pubsub topic
    db = firestore.Client()

    message_data = json.loads(message)
    email = message_data.get('email')
    audio_file_name = message_data.get('audio_file_name')
    dynamic_fields = message_data.get('dynamic_fields', {})
    categories = list(dynamic_fields.values())
    model_type = message_data.get('model_type')
    order_id = message_data.get('order_id')
    amount = message_data.get('amount')
    natural_language_initiator = message_data.get('natural_language_initiator')
    order_ref = db.collection('orders').document(order_id)

    if not check_file_upload(CLOUD_BUCKET_NAME, audio_file_name):
        return "File not uploaded", 400

    task_count = 0
    split_audio_files = split_audio_file_gcs(CLOUD_BUCKET_NAME, audio_file_name)
    task_count = len(split_audio_files)
    for index, split_audio_file in enumerate(split_audio_files, start=1):
        transcript_data = {
            'email': email,
            'audio_file_name': audio_file_name,
            'split_audio_file': split_audio_file,
            'dynamic_fields': dynamic_fields,
            'categories': categories,
            'model_type': model_type,
            'order_id': order_id,
            'amount': amount,
            'natural_language_initiator': natural_language_initiator,
            'order_ref': order_ref.path,
            'index': index
        }
        create_cloud_task(transcript_data)
        
    order_ref.update({'task_count': task_count, 'transcripts_processed': 0})

    return "Tasks created", 200

def create_cloud_task(transcript_data):
#creates a cloud task for each transcript part for the faster processing of the audio files
    client = tasks_v2.CloudTasksClient()
    parent = client.queue_path(project, location, queue)

    url = 'https://worker-service-wtajjbsheq-ez.a.run.app/process_transcript_task'
    payload = json.dumps(transcript_data).encode()

    task = {
        'http_request': {
            'http_method': tasks_v2.HttpMethod.POST,
            'url': url,
            'headers': {'Content-type': 'application/json'},
            'body': payload
        }
    }

    response = client.create_task(request={"parent": parent, "task": task})
    print('Task created: {}'.format(response.name))

def process_transcript(email, audio_file_name, split_audio_file, dynamic_fields, natural_language_initiator, categories, model_type, order_id, index, amount):
#firstly transcribe the audio file, then summarize it if the user has selected the natural language initiator, then create the notule with chatgpt and send it to firestore
    transcript_part = transcribe_audio_gcs(CLOUD_BUCKET_NAME, split_audio_file)

    if natural_language_initiator:
        shorted_full_transcript = summarize(transcript_part, 10)
        topics = extract_topics_with_gpt4(shorted_full_transcript)
        categories += [topic.strip() for topic in topics.split(',')]
    notule_part = create_notule(transcript_part, categories, model_type)

    actionlist_part = create_actionlist(transcript_part, categories)

    order_data = {
        'email': email,
        'audio_file_name': audio_file_name,
        'transcript_part': transcript_part,
        'notule_part': notule_part,
        'actionlist_part': actionlist_part,
        'dynamic_fields': dynamic_fields,
        'natural_language_initiator': natural_language_initiator,
        'model_type': model_type,
        'order_id': order_id,
        'index': index,
        'categories': categories,
        'amount': amount
    }
    
    save_order_data(order_id, index, order_data)
    
    update_transcripts_processed(order_id)
    
    return notule_part

def save_order_data(order_id, index, order_data):
#save the order data in firestore so that we can retrieve it later
    db = firestore.Client()

    order_ref = db.collection('orders').document(order_id)
    general_order_data = {
        'email': order_data['email'],
        'audio_file_name': order_data['audio_file_name'],
        'dynamic_fields': order_data['dynamic_fields'],
        'natural_language_initiator': order_data['natural_language_initiator'],
        'model_type': order_data['model_type'],
        'categories': order_data['categories']
    }
    order_ref.set(general_order_data, merge=True)

    part_data = {
        'transcript_part': order_data['transcript_part'],
        'notule_part': order_data['notule_part'],
        'actionlist_part': order_data['actionlist_part']
    }
    summary_ref = order_ref.collection('summaries').document(str(index))
    summary_ref.set(part_data)

    print(f"Transcript deel {index} en notule deel {index} opgeslagen onder order {order_id}.")

def update_transcripts_processed(order_id):
#update the transcripts processed in firestore so that we can check if all the transcripts have been processed
    db = firestore.Client()
    order_ref = db.collection('orders').document(order_id)

    @firestore.transactional
    def update_in_transaction(transaction, order_ref):
        snapshot = order_ref.get(transaction=transaction)
        if snapshot.exists:
            new_transcripts_processed = snapshot.get('transcripts_processed') + 1
            transaction.update(order_ref, {
                'transcripts_processed': new_transcripts_processed
            })
    
    transaction = db.transaction()
    print(f"transaction in functie update_transcripts_processed {transaction}")
    update_in_transaction(transaction, order_ref)

def summarize(text, per):
#summarize the text with spacy because of old chatgpt input limits and cheaper < might as well replace this in the near future for a better summarizer but i believe this works for now
    doc = nlp(text)
    tokens =[token.text for token in doc]
    word_frequencies={}
    for word in doc:
        if word.text.lower() not in list(STOP_WORDS):
            if word.text.lower() not in punctuation:
                if word.text not in word_frequencies.keys():
                    word_frequencies[word.text] = 1
                else:
                    word_frequencies[word.text] += 1
    max_frequency=max(word_frequencies.values())
    for word in word_frequencies.keys():
        word_frequencies[word]=word_frequencies[word]/max_frequency
    sentence_tokens= [sent for sent in doc.sents]
    sentence_scores = {}
    for sent in sentence_tokens:
        for word in sent:
            if word.text.lower() in word_frequencies.keys():
                if sent not in sentence_scores.keys():                            
                    sentence_scores[sent]=word_frequencies[word.text.lower()]
                else:
                    sentence_scores[sent]+=word_frequencies[word.text.lower()]
    select_length = max(1, int(len(sentence_tokens) * per))
    summary = nlargest(select_length, sentence_scores, key=sentence_scores.get)
    ordered_summary = sorted(summary, key=lambda s: sentence_tokens.index(s))
    final_summary = ' '.join([sent.text for sent in ordered_summary])
    print(final_summary)
    return final_summary

def extract_topics_with_gpt4(text):
#extract topics from the text with gpt4 that will be used to instruct chatgpt how to create the notule when the user has selected the natural language initiator a.k.a. slimme termen detectie
    llm = ChatOpenAI(temperature=0, openai_api_key=os.getenv("OPENAI_API_KEY"), model_name="gpt-4-1106-preview")

    template ="""
        Identificeer de grotere onderwerpen/thema's in de volgende tekst. 
        Deze zullen als kopjes voor in de notule gebruikt worden, benoem echt alleen de thema's gescheiden door een kommateken.
        Dit is de tekst waarbij dat moet gebeuren:
        "{docs}"
        """
    prompt = PromptTemplate(
        template=template, input_variables=["docs"]
    )

    llmchain = LLMChain(
        llm=llm,
        prompt=prompt,
        )
    
    llm_output = llmchain.predict(docs=text)

    print (f"Dit is de LLM_output in extract_topics_with_gpt4 {llm_output}!!")
    return llm_output

def format_summary_text(text):
#format text after it has been retrieved from firestore to make it more readable, this outputs like:
# Subject
# - text for subject
# - another text for subject
#
# Subject 2
# - text for subject 2
# - another text for subject 2
    formatted_lines = []
    for line in text.split('\n'):
        if line.endswith(':'):
            formatted_lines.append("\n" + line)
        else:
            formatted_lines.extend([f"- {x.strip()}" for x in line.split('-') if x.strip()])
    formatted_text = '\n'.join(formatted_lines)
    return formatted_text

def check_file_upload(bucket_name, blob_name, retries=3, delay=10):
#check if file is uploaded
    storage_client = storage.Client()
    bucket = storage_client.bucket(bucket_name)
    
    for i in range(retries):
        blob = bucket.get_blob(blob_name)
        if blob and blob.size:
            print(f"Bestand {blob_name} is succesvol geüpload met grootte {blob.size} bytes.")
            return True
        print(f"Bestand {blob_name} is nog niet geüpload, wachten {delay} seconden...")
        time.sleep(delay)
        delay *= 2
    
    print(f"Bestand {blob_name} is niet geüpload na {retries} pogingen.")
    return False

def get_file_size_from_gcs(bucket_name, blob_name):
#check for filesize in gcs, again because of the whisper limit and the other limit that i cant remember
    storage_client = storage.Client()
    bucket = storage_client.bucket(bucket_name)
    blob = bucket.get_blob(blob_name)
    return blob.size if blob else 0

def transcribe_audio_gcs(bucket_name, blob_name):
#transcribe the actual audio with OpenAi whisper
    storage_client = storage.Client()
    bucket = storage_client.bucket(bucket_name)
    blob = bucket.blob(blob_name)
    
    audio_data = blob.download_as_bytes()
    audio_file = BytesIO(audio_data)
    audio_file.name = blob_name

    client = OpenAI()

    transcript = client.audio.transcriptions.create(
        model="whisper-1", 
        file=audio_file,
        response_format="text"
    )
    
    return transcript

def split_audio_file_gcs(bucket_name, blob_name):
#split the audio file into roughly 10 minute chunks, because of the 20mb limit of openai whisper and some other limit that i can't remember...
    storage_client = storage.Client()
    bucket = storage_client.bucket(bucket_name)
    blob = bucket.blob(blob_name)
    
    audio_data = blob.download_as_bytes()
    audio = AudioSegment.from_file(io.BytesIO(audio_data), format="mp3")
    
    chunk_length = 450000
    chunks = [audio[i:i + chunk_length] for i in range(0, len(audio), chunk_length)]
    
    split_files = []
    for i, chunk in enumerate(chunks):
        buffer = io.BytesIO()
        chunk.export(buffer, format="mp3")
        
        buffer_size = len(buffer.getbuffer())
        if buffer_size > 20 * 1024 * 1024:
            continue
        
        split_file_path = f"{blob_name}_part_{i}.mp3"
        split_blob = bucket.blob(split_file_path)
        split_blob.upload_from_string(buffer.getvalue(), content_type='audio/mp3')
        split_files.append(split_file_path)

    return split_files

def create_notule(transcript, categories, model_type):
#create the notule with chatgpt and with the categories
    dynamic_prompt_parts = ["Jij bent een notulist die een vergadering moet samenvatten. \nHou rekening met de volgende instructies: \n"]
    for cat in categories:
        dynamic_prompt_parts.append(f"{cat}:")
    dynamic_prompt_parts.append("'")
    dynamic_prompt = "\n".join(dynamic_prompt_parts)
    if model_type == 'gpt3':
        llm = ChatOpenAI(temperature=0, openai_api_key=os.getenv("OPENAI_API_KEY"), model_name="gpt-3.5-turbo-1106")
    elif model_type == 'gpt4':
        llm = ChatOpenAI(temperature=0, openai_api_key=os.getenv("OPENAI_API_KEY"), model_name="gpt-4-1106-preview")

    
    template ="""
    Instructies voor het verwerken van de vergadering:
    - Ga zorgvuldig door het onderstaande deel van de vergadering.
    - Jouw taak is om een gedetailleerde en grondige samenvatting te geven onder elk van de volgende kopjes. De samenvatting moet voldoende informatie bevatten zodat iemand die niet aanwezig was bij de vergadering een duidelijk en compleet beeld krijgt van de besproken punten.
    - Gebruik opsommingstekens ("-") voor elk punt dat je samenvat. Zorg ervoor dat elk punt vol staat met informatie en context. 
    - VERMIJD het direct citeren van de vergadering. Parafreseer in plaats daarvan de informatie om diepte en duidelijkheid te geven.
    - Let op de details: wie zegt wat, welke besluiten worden genomen en welke nuances in de discussie worden uitgesproken.
    - Maak geen nieuwe kopjes aan, gebruik alleen de kopjes die hieronder staan.
    - Wanneer iets besproken, verwezen of gezegd wordt, zorg er dan voor dat je ook die informatie samenvat in het punt.

    GEGEVEN KOPJES VOOR EXTRACTIE:
    {dynamic_prompt}

    DEEL VAN DE VERGADERING:
    {docs}
    """
    prompt = PromptTemplate(
        template=template, input_variables=["dynamic_prompt", "docs"]
    )

    llmchain = LLMChain(
        llm=llm,
        prompt=prompt,
        )
    
    results = llmchain.predict(dynamic_prompt=dynamic_prompt, docs=transcript)
    
    return results

def create_actionlist(transcript, categories):
#create actionlist from the transcripts with chatgpt
    dynamic_prompt_parts = ["**Instructies:**\n"]
    for cat in categories:
        dynamic_prompt_parts.append(f"{cat}:")
    dynamic_prompt_parts.append("'")
    dynamic_prompt = "\n".join(dynamic_prompt_parts)
    llm = ChatOpenAI(temperature=0, openai_api_key=os.getenv("OPENAI_API_KEY"), model_name="gpt-4-1106-preview")
    
    template ="""
    Gebruik de volgende tekst om alle specifieke actiepunten en de eventuele namen van de personen die verantwoordelijk zijn voor deze acties te identificeren. 
    Zorg ervoor dat van elk actiepunt duidelijk is wat er moet gebeuren, ook voor mensen die niet aanwezig waren, geef dus genoeg context.
    Nummer de actiepunten NIET.
    Voeg ook ALTIJD een witregel na de laatste actiepunt toe.
    Gebruik altijd dit format:
    **Format:**
    -Actiepunt: Voorbeeld actiepunt
    -Onderwerp: Voorbeeld onderwerp
    -Context: Voorbeeld context
    -Verantwoordelijke: Voorbeeld eventuele verantwoordelijke

    Focus alleen op concrete taken en verantwoordelijke personen, en negeer alle overige informatie. 
    Gebruik deze lijst met Onderwerpen om de onderwerpen te identificeren:
    **Onderwerpen:**
    {dynamic_prompt}

    Hieronder staat de tekst waar de actiepunten uitgehaald moeten worden:
    **Tekst:**
    {docs}
    """
    prompt = PromptTemplate(
        template=template, input_variables=["dynamic_prompt", "docs"]
    )

    llmchain = LLMChain(
        llm=llm,
        prompt=prompt,
        )
    
    results = llmchain.predict(dynamic_prompt=dynamic_prompt, docs=transcript)
    
    return results

def create_summary_dict(formatted_summary):
#create a dictionary from the formatted summary text
    summary_dict = {}
    current_topic = None
    for line in formatted_summary.split('\n'):
        if line.startswith('-'):
            if current_topic:
                summary_dict[current_topic] += line + '\n'
        else:
            current_topic = line.strip(':')
            summary_dict[current_topic] = ''
    return summary_dict

def combine_summaries(organized_output, summary_dict):
#combine the summaries with the formatted summary text and the summary dictionary, this together with create_summary_dict is used to make text output like this:
# Overarching topic
#     Smaller subject from the overarching topic
#     - text for smaller subject
#     - another text for smaller subject
#    
#     Smaller subject 2 from the overarching topic
#     - text for smaller subject 2
#     - another text for smaller subject 2
    combined_output = ''
    for line in organized_output.split('\n'):
        if line.startswith('?'):
            # Nieuw overkoepelend thema in hoofdletters
            current_topic = line[1:].strip().upper()
            combined_output += current_topic + '\n'
        elif line.startswith('-'):
            # Klein onderwerp onder het huidige thema
            sub_topic = line[2:].strip()
            combined_output += '  ' + sub_topic + '\n'
            if sub_topic in summary_dict:
                # Voeg de bijbehorende tekst toe voor dit kleine onderwerp
                text_lines = summary_dict[sub_topic].split('\n')
                for text_line in text_lines:
                    # Controleer of de tekstregel niet leeg is
                    if text_line.strip():
                        # Verwijder het eerste streepje als het er is
                        if text_line.startswith('-'):
                            text_line = text_line[1:].strip()
                        combined_output += '    - ' + text_line + '\n'
        else:
            combined_output += line + '\n'
    return combined_output

def organiseer_tekst_per_categorie(categories, text):
    categorie_data = {categorie: [] for categorie in categories}
    huidige_categorie = None

    for regel in text.split('\n'):
        if any(categorie in regel for categorie in categories):
            huidige_categorie = next((categorie for categorie in categories if categorie in regel), None)
        
        if huidige_categorie and regel not in categorie_data[huidige_categorie]:
            categorie_data[huidige_categorie].append(regel)

    llm = ChatOpenAI(temperature=0, openai_api_key=os.getenv("OPENAI_API_KEY"), model_name="gpt-4-1106-preview")
    
    finale_output = ""
    for categorie in categories:
        if categorie_data[categorie]:
            categorie_tekst = '\n'.join(categorie_data[categorie])
            # Gebruik GPT-4 om de tekst per categorie na te kijken
            prompt = f"""Haal de dubbele of lege inputs uit de tekst weg, zodat het een meer waardevolle en leesbare tekst wordt.
            Haal wel altijd het koptekst/de categorie aan met daaronder de regels aangeduidt met een opsommingsteken "-"
            
            Dit is de tekst die je moet nakijken:
            {categorie_tekst}"""
            llm_output = llm.predict(prompt)
            # Voeg GPT-4 output toe aan de finale output, met opsommingstekens
            geformatteerde_output = "\n".join(f"- {line}" for line in llm_output.split('\n'))
            finale_output += f"{categorie}:\n{geformatteerde_output}\n\n"

    return finale_output

def parse_action_list(action_list_text):
# Functie om de actielijst te parsen en om te zetten in een lijst van dictionaries
    lines = action_list_text.split('\n')
    action_items = []
    temp_item = {}
    for line in lines:
        if '-Actiepunt:' in line:
            if temp_item:
                action_items.append(temp_item)
                temp_item = {}
            temp_item['Actiepunt'] = line.split('-Actiepunt:', 1)[-1].strip()
        elif '-Onderwerp:' in line:
            temp_item['Onderwerp'] = line.split('-Onderwerp:', 1)[-1].strip()
        elif '-Context:' in line:
            temp_item['Context'] = line.split('-Context:', 1)[-1].strip()
        elif '-Verantwoordelijke:' in line:
            temp_item['Verantwoordelijke'] = line.split('-Verantwoordelijke:', 1)[-1].strip()

    # Voeg het laatste item toe als het niet leeg is
    if temp_item:
        action_items.append(temp_item)

    return action_items

def add_action_points_table(doc, action_list):
# Functie om een tabel toe te voegen aan het Word-document voor de actielijst
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells

    hdr_cells[0].text = 'Actiepunt'
    hdr_cells[0].paragraphs[0].runs[0].bold = True

    hdr_cells[1].text = 'Onderwerp'
    hdr_cells[1].paragraphs[0].runs[0].bold = True

    hdr_cells[2].text = 'Context'
    hdr_cells[2].paragraphs[0].runs[0].bold = True

    hdr_cells[3].text = 'Verantwoordelijke'
    hdr_cells[3].paragraphs[0].runs[0].bold = True

    for item in action_list:
        row_cells = table.add_row().cells
        row_cells[0].text = item['Actiepunt']
        row_cells[1].text = item['Onderwerp']
        row_cells[2].text = item['Context']
        row_cells[3].text = item['Verantwoordelijke']

def create_word_document(notule, notule_detailed, finalaction_list, transcript_with_speakers, transcript):
    doc = Document()

    def add_section(header, content, is_table=False):
        doc.add_heading(header, level=1)
        if not is_table:
            doc.add_paragraph(content)
        else:
            action_list = parse_action_list(content)
            add_action_points_table(doc, action_list)
        doc.add_page_break()

    add_section("Notulen", notule)
    add_section("Actielijst", finalaction_list, is_table=True)
    add_section("Notulen in Details", notule_detailed)
    add_section("Volledige Transcriptie met Sprekers", transcript_with_speakers)
    add_section("Volledige Transcriptie Whisper", transcript)

    doc.save("vergadering_documentatie.docx")

def send_email_final(email, notule, notule_detailed, finalaction_list, transcript_with_speakers, transcript, user_friendly_data):
    from_email = os.environ.get("FROM_EMAIL")
    email_password = os.environ.get("EMAIL_PASSWORD")
    to_email = email

    # Combineer alle inhoud in één Word-document
    create_word_document(notule, notule_detailed, finalaction_list, transcript_with_speakers, transcript)

    # E-mail opstellen
    subject = "Uw vergadertranscript en samenvatting"
    msg = MIMEMultipart()
    msg['From'] = from_email
    msg['To'] = to_email
    msg['Subject'] = subject

    body = f'''
        <html>
            <head></head>
            <body>
                <p>Hierbij ontvangt u de volledige documentatie van uw vergadering in één Word-bestand.</p>
                <p>Details van uw order:</p>
                <br>
                <ul>
                    <li><strong>Order ID:</strong> {user_friendly_data['Order ID']}</li>
                    <li><strong>Bestandsnaam:</strong> {user_friendly_data['Bestandsnaam']}</li>
                    <li><strong>E-mail:</strong> {user_friendly_data['E-mail']}</li>
                    <li><strong>Model Type:</strong> {user_friendly_data['Model Type']}</li>
                    <li><strong>Onderwerpen:</strong> {user_friendly_data['Onderwerpen']}</li>
                    <li><strong>Prijs (in centen):</strong> {user_friendly_data['Prijs (in centen)']}</li> 
                </ul>
                <br>
                <p><bold><italic>Groetjes, NOA ;)</p></bold></italic>
            </body>
        </html>
        '''
    msg.attach(MIMEText(body, 'html'))

    # Voeg het Word-document toe als bijlage
    filename = "vergadering_documentatie.docx"
    attachment = open(filename, "rb")
    part = MIMEBase('application', 'octet-stream')
    part.set_payload((attachment).read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', "attachment; filename= %s" % filename)
    msg.attach(part)

    # E-mail versturen
    smtp_server = "mail.privateemail.com"
    smtp_port = 587
    server = smtplib.SMTP(smtp_server, smtp_port)
    server.starttls()
    server.login(from_email, email_password)
    server.sendmail(from_email, to_email, msg.as_string())
    server.quit()

def send_email(email, user_friendly_data):
#send the email with the transcript and summary to the user and update firestore to say that the email has been sent
    from_email = os.environ.get("FROM_EMAIL")
    email_password = os.environ.get("EMAIL_PASSWORD")
    to_email = email

    subject = "Verwerking is gestart!"

    msg = MIMEMultipart()
    msg['From'] = from_email
    msg['To'] = to_email
    msg['Subject'] = subject

    body = f'''
        <html>
            <head></head>
            <body>
                <p>Uw bestand is succesvol geüpload en wordt nu verwerkt.</p>
                <p>Details van uw order:</p>
                <br>
                <ul>
                    <li><strong>Order ID:</strong> {user_friendly_data['Order ID']}</li>
                    <li><strong>Bestandsnaam:</strong> {user_friendly_data['Bestandsnaam']}</li>
                    <li><strong>E-mail:</strong> {user_friendly_data['E-mail']}</li>
                    <li><strong>Model Type:</strong> {user_friendly_data['Model Type']}</li>
                    <li><strong>Onderwerpen:</strong> {user_friendly_data['Onderwerpen']}</li>
                    <li><strong>Prijs (in centen):</strong> {user_friendly_data['Prijs (in centen)']}</li> 
                </ul>
                <br>
                <p><bold><italic>Groetjes, NOA ;)</p></bold></italic>
            </body>
        </html>
        '''
    msg.attach(MIMEText(body, 'html'))

    smtp_server = "mail.privateemail.com"
    smtp_port = 587 

    server = smtplib.SMTP(smtp_server, smtp_port)
    server.starttls()
    server.login(from_email, email_password)
    text = msg.as_string()
    server.sendmail(from_email, to_email, text)
    server.quit()

def create_actual_cloud_task_refine(data):
    client = tasks_v2.CloudTasksClient()
    parent = client.queue_path(project, location, queue)

    url = 'https://worker-service-wtajjbsheq-ez.a.run.app/process_part'
    payload = json.dumps(data).encode()

    task = {
        'http_request': {
            'http_method': tasks_v2.HttpMethod.POST,
            'url': url,
            'headers': {'Content-type': 'application/json'},
            'body': payload
        }
    }

    response = client.create_task(request={"parent": parent, "task": task})
    print('Task created: {}'.format(response.name))

def save_order_data_refine(order_id, index, order_data):
#save the order data in firestore so that we can retrieve it later
    db = firestore.Client()

    order_ref = db.collection('orders').document(order_id)

    part_data = {
        'transcript_with_speakers': order_data['transcript_with_speakers'],
    }
    summary_ref = order_ref.collection('summaries').document(str(index))
    summary_ref.update(part_data)

    print(f"Notule_refine deel {index} opgeslagen onder order {order_id}.")

def update_transcripts_processed_refine(order_id):
#update the transcripts processed in firestore so that we can check if all the transcripts have been processed
    db = firestore.Client()
    order_ref = db.collection('orders').document(order_id)

    @firestore.transactional
    def update_in_transaction(transaction, order_ref):
        snapshot = order_ref.get(transaction=transaction)
        if snapshot.exists:
            new_transcripts_processed = snapshot.get('processed_parts_refine') + 1
            transaction.update(order_ref, {
                'processed_parts_refine': new_transcripts_processed
            })
    
    transaction = db.transaction()
    print(f"transaction in functie update_transcripts_processed_refine {transaction}")
    update_in_transaction(transaction, order_ref)

def process_transcript_refine(notule_part, order_id, index):
    qdrant = QdrantClient(url=os.getenv("QDRANT_URL"), api_key=os.getenv("QDRANT_API_KEY"))
    embeddings = OpenAIEmbeddings()
    llm = ChatOpenAI(temperature=0, openai_api_key=os.getenv("OPENAI_API_KEY"), model_name="gpt-4-1106-preview")
    vector_store = Qdrant(client=qdrant, collection_name=order_id, embeddings=embeddings)
    qa = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=vector_store.as_retriever())

    if not notule_part:
        logging.info(f"Geen notule deel gevonden voor order_id {order_id}")
        return jsonify({"error": "Geen notule deel gevonden"}), 400
    query = f"""
    Zoek de sprekers bij de gegeven tekst en maak een samenvatting van de regels onder het THEMA, in BOLD/dikgedrukt.
    Zoek ook aanvullende informatie bij abstracte zinnen.
    Zoek ook of de gegeven regels kloppen in de context voor je deze samenvat.

    Dit is het format wat je moet aanhouden om de gegeven tekst daarna mee terug te geven, zet voor het THEMA/ONDERWERP altijd een witregel:

    THEMA
    Daaronder een samenvatting van de regels onder het thema met de sprekers en abstracte zinnen verduidelijkt


    Gegeven tekst dat verrijkt moet worden met sprekers en correctie met meer duidelijkere taal:
    {notule_part}"""

    # Voer de query uit
    transcript_with_speakers = qa.run(query)

    order_data = {
        'transcript_with_speakers': transcript_with_speakers,
        'order_id': order_id,
        'index': index,
    }
    
    save_order_data_refine(order_id, index, order_data)
    
    update_transcripts_processed_refine(order_id)
    
    return transcript_with_speakers

@app.route('/process_message', methods=['POST'])
@cross_origin()
def process_message_endpoint():
#endpoint for the pubsub topic to send messages to the worker service to process the audio file
    message_data = request.json
    app.logger.info(f"Received message data: {message_data}")

    if not message_data:
        return jsonify({"error": "No message_data provided"}), 400

    encoded_data = message_data.get('message', {}).get('data')
    if encoded_data:
        decoded_data = base64.b64decode(encoded_data).decode('utf-8')
        message_data_json = json.loads(decoded_data)

        email = message_data_json.get('email')
        audio_file_name = message_data_json.get('audio_file_name')
    else:
        email = None
        audio_file_name = None

    app.logger.info(f"Email: {email}, Audio File Name: {audio_file_name}")

    if not email or not audio_file_name:
        return jsonify({"error": "Missing email or audio_file_name"}), 400

    url = 'https://worker-service-wtajjbsheq-ez.a.run.app/task_handler'
    payload = json.dumps(message_data_json).encode()

    task = {
        'http_request': {
            'http_method': 'POST',
            'url': url,
            'headers': {
                'Content-type': 'application/json',
            },
            'body': payload
        }
    }

    client.create_task(request={"parent": parent, "task": task})

    return "Task queued", 200

@app.route('/task_handler', methods=['POST'])
def task_handler():
#handler for the cloud task that is created in the process_message_endpoint
    payload = json.loads(request.data.decode())
    process_message(json.dumps(payload))
    return "Task completed", 200

@app.route('/process_transcript_task', methods=['POST'])
def process_transcript_task():
#route to create the task for the transcript processing
    task_params = request.json

    result = process_transcript(
        task_params['email'],
        task_params['audio_file_name'],
        task_params['split_audio_file'],
        task_params['dynamic_fields'],
        task_params['natural_language_initiator'],
        task_params['categories'],
        task_params['model_type'],
        task_params['order_id'],
        task_params['index'],
        task_params['amount']
    )

    return jsonify({"result": "success", "summary": result}), 200

@app.route('/process_final_results', methods=['POST'])
def process_final_results():
#the final step to bring all the processed texts together and create the final summary
    order_id = request.args.get('order_id')
    
    if not order_id:
        return jsonify({"error": "Order ID not provided"}), 400

    db = firestore.Client()
    order_ref = db.collection('orders').document(order_id)
    order_doc = order_ref.get()

    if not order_doc.exists:
        return jsonify({"error": "Order not found"}), 404

    order_data = order_doc.to_dict()
    if order_data.get('email_sent', False):
        return jsonify({"message": "E-mail is al verzonden"}), 200 
    
    email = order_data.get('email') 
    amount = order_data.get('amount')
    audio_file_name = order_data.get('audio_file_name')
    dynamic_fields = order_data.get('dynamic_fields')
    natural_language_initiator = order_data.get('natural_language_initiator')
    model_type = order_data.get('model_type')
    categories = order_data.get('categories')

    user_friendly_data = {
        'Bestandsnaam': audio_file_name,
        'E-mail': email,
        'Model Type': 'GPT-3' if model_type == 'gpt3' else 'GPT-4',
        'Onderwerpen': categories,
        'Order ID': order_id,
        'Prijs (in centen)': amount
        }

    send_email(email, user_friendly_data)

    summaries_ref = order_ref.collection('summaries')
    summaries_docs = summaries_ref.stream()

    sorted_summaries = sorted(
        summaries_docs,
        key=lambda doc: int(doc.id)
    )

    sorted_summaries_dicts = [doc.to_dict() for doc in sorted_summaries]

    final_transcript = ""
    final_summary = ""
    final_actionlist = ""

    for doc in sorted_summaries_dicts:
        transcript_part = doc.get('transcript_part', '')
        notule_part = doc.get('notule_part', '')
        actionlist_part = doc.get('actionlist_part', '')

        final_transcript += transcript_part + "\n"
        final_summary += notule_part + "\n"
        final_actionlist += actionlist_part + "\n"

    final_transcript = final_transcript.strip()
    final_summary = final_summary.strip()
    final_actionlist = final_actionlist.strip()

    formatted_final_summary = format_summary_text(final_summary)

    llm = ChatOpenAI(temperature=0, openai_api_key=os.getenv("OPENAI_API_KEY"), model_name="gpt-4-1106-preview")

    template ="""
        Gebruik uit de volgende tekst de onderwerpen welke gebruikt worden als kopjes en kijk of je hier samenhangende grotere overkoepelende onderwerpen/thema's van kan maken.
        Deze zullen worden gebruikt als kopjes voor boven de onderwerpen in de notule.
        Het moeten dan dus ook meerdere grotere overkoepelende onderwerpen/thema's worden, maar nooit net zo veel als de onderwerpen zelf.
        Geef vervolgens de gebruikte onderwerpen terug onder het nieuwe overkoepelende onderwerp.
        Scheidt de gebruikte onderwerpen door een "!" ervoor te zetten en zet een "?" voor het nieuwe overkoepelende onderwerp.
        Scheidt deze setjes door een witregel.
        Hier is de volledige tekst:
        "{docs}"
        """
    prompt = PromptTemplate(
        template=template, input_variables=["docs"]
    )

    llmchain = LLMChain(
        llm=llm,
        prompt=prompt,
        )
    
    llm_output = ""
    organized_output = ""
    
    if natural_language_initiator:
        llm_output = llmchain.predict(docs=formatted_final_summary)

        current_topic = None
        for line in llm_output.split('\n'):
            if line.startswith('?'):
                current_topic = line[1:].strip().upper()
                organized_output += f"\n{current_topic}\n"
            elif line.startswith('!'):
                sub_topics = line[1:].split('!')
                for sub_topic in sub_topics:
                    if sub_topic.strip():
                        organized_output += f"- {sub_topic.strip()}\n"
        
        summary_dict = create_summary_dict(formatted_final_summary)
        final_combined_output_notule = combine_summaries(organized_output, summary_dict)
    else:
        final_combined_output_notule = organiseer_tekst_per_categorie(categories, formatted_final_summary)

    extra_data = {
    'final_transcript': final_transcript,
    'final_actionlist': final_actionlist,
    'final_combined_output': final_combined_output_notule,
    }
    order_ref.update(extra_data)

    order_data = {
            'transcriptie_whisper': 'klaar',
        }
    order_ref.update(order_data)

    return "Message processed", 200

@app.route('/combine_transcript_with_speakers', methods=['POST'])
def combine_transcript_with_speakers():
    order_id = request.args.get('order_id')
    db = firestore.Client()
    order_ref = db.collection('orders').document(order_id)

    summary_ref = order_ref.collection('summaries')
    notule_parts = [doc.to_dict() for doc in summary_ref.stream()]

    total_tasks = 0
    total_tasks = len(notule_parts)

    for index, notule_part in enumerate(notule_parts, start=1):
        data = {
            "order_id": order_id,
            "notule_part": notule_part,
            "index": index,
            'order_ref': order_ref.path,
        }

        create_actual_cloud_task_refine(data)

    order_ref.update({'processing_started': True, 'total_tasks_refine': total_tasks, 'processed_parts_refine': 0})
    
    return "Tasks created", 200

@app.route('/process_part', methods=['POST'])
def process_part():
    #route to create the task for the transcript refining processing
    task_params = request.json

    result = process_transcript_refine(
        task_params['notule_part'],
        task_params['order_id'],
        task_params['index'],
    )

    return jsonify({"result": "success", "summary": result}), 200

@app.route('/final_route', methods=['POST'])
def final_route():
    order_id = request.args.get('order_id')

    # Haal de ordergegevens op
    db = firestore.Client()
    order_ref = db.collection('orders').document(order_id)
    order_doc = order_ref.get()

    if not order_doc.exists:
        return jsonify({"error": "Order not found"}), 404

    order_data = order_doc.to_dict()

    summaries_ref = order_ref.collection('summaries')
    summaries_docs = summaries_ref.stream()

    sorted_summaries = sorted(
        summaries_docs,
        key=lambda doc: int(doc.id)
    )

    sorted_summaries_dicts = [doc.to_dict() for doc in sorted_summaries]

    final_notule = ""

    for doc in sorted_summaries_dicts:
        notule_refine_part = doc.get('transcript_with_speakers', '')

        final_notule += notule_refine_part + "\n"

    final_notule = final_notule.strip()

    email = order_data.get('email') 
    amount = order_data.get('amount')
    audio_file_name = order_data.get('audio_file_name')
    model_type = order_data.get('model_type')
    categories = order_data.get('categories')
    final_transcript = order_data.get('final_transcript')
    final_combined_output_notule = order_data.get('final_combined_output')
    final_actionlist = order_data.get('final_actionlist')
    full_transcript_with_speakers = order_data.get('full_transcript_with_speakers')
    user_friendly_data = {
        'Bestandsnaam': audio_file_name,
        'E-mail': email,
        'Model Type': 'GPT-3' if model_type == 'gpt3' else 'GPT-4',
        'Onderwerpen': categories,
        'Order ID': order_id,
        'Prijs (in centen)': amount
        }

    send_email_final(email, final_notule, final_combined_output_notule, final_actionlist, full_transcript_with_speakers, final_transcript, user_friendly_data)

    # Update de ordergegevens
    order_ref.update({
        'final_notule': final_notule,
        'email_sent': True,
        'status': "done"
    })

    client = QdrantClient(
        os.getenv("QDRANT_URL"),
        api_key=os.getenv("QDRANT_API_KEY"),
    )
    client.delete_collection(collection_name=order_id)

    return jsonify({"message": "E-mail met gecombineerde notulen verzonden"}), 200

if __name__ == '__main__':
    app.run(port=int(os.environ.get("PORT", 8080)), host='0.0.0.0', debug=True)